#!/usr/bin/env python3
"""
slack_attachment_reader.py

Slackの指定チャンネルから直近のメッセージを取得し、添付されている
Excel (.xlsx / .xls) / Word (.docx) / CSV ファイルの中身を
テキスト化して標準出力 or ファイルに書き出す。

------------------------------------------------------------
セットアップ (User Token方式)
------------------------------------------------------------
1. https://api.slack.com/apps から Slack App を作成
2. "OAuth & Permissions" で User Token Scopes に以下を付与:
        - channels:history      (publicチャンネルの読み取り)
        - channels:read         (チャンネル名→ID解決用)
        - files:read            (添付ファイルのダウンロード)
        - groups:history        (privateチャンネル)
        - groups:read
        - im:history            (DM、必要なら)
        - mpim:history          (グループDM、必要なら)
3. ワークスペースに install して xoxp- で始まる User OAuth Token を取得
4. 依存ライブラリをインストール:
        pip install slack_sdk requests openpyxl python-docx pandas xlrd \
                    python-dotenv msal
5. 環境変数にトークンを設定 (or .envファイル):
        SLACK_TOKEN=xoxp-xxxxxxxx...

SharePoint連携 (任意):
- Slackメッセージ内のSharePoint URLからファイルをDLしたい場合は msal が必要
- 初回実行時にデバイスコードフロー認証 (ブラウザでログイン)
- 認証情報は ~/.slack_attachment_reader_msal_cache.bin にキャッシュされ、
  ~90日間は自動更新でログイン不要

------------------------------------------------------------
使い方
------------------------------------------------------------
    # 直近20件のメッセージから添付を読み出す
    python slack_attachment_reader.py --channel C0123456789

    # チャンネル名でも指定可。スレッド内の返信も対象にする
    python slack_attachment_reader.py --channel '#general' --threads

    # 件数と期間を指定し、結果をファイルへ
    python slack_attachment_reader.py --channel sales --limit 100 \
        --since 2026-04-01 --out result.txt
"""

import argparse
import io
import os
import re
import sys
from datetime import date, datetime
from pathlib import Path

import requests
from slack_sdk import WebClient
from slack_sdk.errors import SlackApiError

# .env ファイルがあれば読み込む (任意、無くても動く)
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass


TARGET_EXTS = {".xlsx", ".xls", ".docx", ".csv"}

# 例: https://workspace.slack.com/files/U123/F0ABC123/filename.xlsx
SLACK_FILE_ID_RE = re.compile(r'/files/[A-Z0-9]+/(F[A-Z0-9]+)')

# 例: https://mycompany.sharepoint.com/sites/.../運行20260520.xlsx
SHAREPOINT_URL_RE = re.compile(
    r'https://[\w-]+\.sharepoint\.com/[^\s<>"\'|]+',
    re.IGNORECASE,
)

# 例: https://teams.microsoft.com/l/meetup-join/19%3ameeting_XXX...
TEAMS_URL_RE = re.compile(
    r'https://teams\.microsoft\.com/l/meetup-join/[^\s<>"\'|]+',
    re.IGNORECASE,
)


def extract_teams_urls(msg: dict) -> list[str]:
    """メッセージから Teams 会議URL (https://teams.microsoft.com/l/meetup-join/...) を抽出"""
    urls: set[str] = set()
    text = msg.get("text", "") or ""
    urls.update(TEAMS_URL_RE.findall(text))
    for att in msg.get("attachments", []) or []:
        for key in ("from_url", "original_url", "title_link", "fallback", "text"):
            v = att.get(key) or ""
            urls.update(TEAMS_URL_RE.findall(v))
    return list(urls)

# Microsoft Graph CLI Tools の公式アプリID (Azure ADへの登録不要)
GRAPH_CLIENT_ID = "14d82eec-204b-4c2f-b7e8-296a70dab67e"
MSAL_CACHE_FILE = os.path.expanduser("~/.slack_attachment_reader_msal_cache.bin")


def extract_sharepoint_urls(msg: dict) -> list[str]:
    """メッセージ本文・attachments・blocks から SharePoint URLを抽出"""
    urls: set[str] = set()
    text = msg.get("text", "") or ""
    urls.update(SHAREPOINT_URL_RE.findall(text))
    for att in msg.get("attachments", []) or []:
        for key in ("from_url", "original_url", "title_link", "fallback", "text"):
            v = att.get(key) or ""
            urls.update(SHAREPOINT_URL_RE.findall(v))
    return list(urls)


def filename_from_sharepoint_url(url: str) -> str:
    """SharePoint URLからファイル名(末尾セグメント)を推定。URLデコードする。"""
    from urllib.parse import unquote, urlparse
    path = urlparse(url).path
    last = path.rstrip("/").split("/")[-1]
    return unquote(last)


def extract_referenced_file_ids(msg: dict) -> list[str]:
    """メッセージ本文・attachments・blocks から Slack ファイルID (F...) を抽出"""
    ids = []
    text = msg.get("text", "") or ""
    ids.extend(SLACK_FILE_ID_RE.findall(text))

    for att in msg.get("attachments", []) or []:
        if att.get("file_id"):
            ids.append(att["file_id"])
        for key in ("from_url", "original_url", "title_link", "image_url"):
            url = att.get(key, "") or ""
            ids.extend(SLACK_FILE_ID_RE.findall(url))

    for block in msg.get("blocks", []) or []:
        if block.get("type") == "file" and block.get("file_id"):
            ids.append(block["file_id"])

    return list(dict.fromkeys(ids))


# ---------- Slack helpers ----------

def get_token() -> str:
    token = os.environ.get("SLACK_TOKEN") or os.environ.get("SLACK_BOT_TOKEN")
    if not token:
        sys.exit("環境変数 SLACK_TOKEN (xoxp-...) が設定されていません。")
    return token


def resolve_channel_id(client: WebClient, channel: str) -> str:
    """チャンネル名("#general"等)が渡されたらIDに変換。C/G/Dで始まるIDはそのまま返す。"""
    channel = channel.lstrip("#")
    if channel[:1] in {"C", "G", "D"} and channel[1:].isalnum() and channel.isupper():
        return channel

    # まず public + private で検索。groups:read が無い場合は public のみで再試行
    for types in ("public_channel,private_channel", "public_channel"):
        cursor = None
        try:
            while True:
                resp = client.conversations_list(limit=200, cursor=cursor, types=types)
                for ch in resp["channels"]:
                    if ch["name"] == channel:
                        return ch["id"]
                cursor = resp.get("response_metadata", {}).get("next_cursor")
                if not cursor:
                    break
            break  # 検索完了 (見つからなかった)
        except SlackApiError as e:
            if e.response.get("error") == "missing_scope" and "private" in types:
                print("[情報] groups:read 未付与のため publicチャンネルのみで再検索します",
                      file=sys.stderr)
                continue
            raise
    sys.exit(f"チャンネル '{channel}' が見つかりません。"
             "(privateチャンネルを読みたい場合は Slack App に groups:read / groups:history を追加してください)")


def fetch_messages(client: WebClient, channel_id: str, limit: int,
                   since_ts: str | None = None, include_threads: bool = False) -> list[dict]:
    """指定チャンネルからメッセージを取得。include_threadsならスレッド返信も含める。"""
    kwargs = {"channel": channel_id, "limit": limit}
    if since_ts:
        kwargs["oldest"] = since_ts
    resp = client.conversations_history(**kwargs)
    messages = list(resp["messages"])

    if include_threads:
        import time
        extra = []
        parents_with_replies = [
            p for p in messages
            if p.get("reply_count", 0) > 0 and p.get("thread_ts")
        ]

        def thread_url(parent_msg: dict) -> str:
            """エラー出力用にスレッドのpermalinkを取得 (失敗時は手組み)"""
            try:
                r = client.chat_getPermalink(
                    channel=channel_id, message_ts=parent_msg["thread_ts"]
                )
                return r.get("permalink", "")
            except SlackApiError:
                # フォールバック: ts と channel_id から手組み
                ts_clean = parent_msg["thread_ts"].replace(".", "")
                return f"(channel={channel_id} ts={parent_msg['thread_ts']} p{ts_clean})"

        for i, parent in enumerate(parents_with_replies):
            try:
                replies = client.conversations_replies(
                    channel=channel_id, ts=parent["thread_ts"], limit=200
                )
                for reply in replies["messages"][1:]:  # 先頭は親なので除外
                    reply["_parent_msg"] = parent
                    extra.append(reply)
            except SlackApiError as e:
                err = e.response.get("error", "")
                if err == "ratelimited":
                    wait = int(e.response.headers.get("Retry-After", "30"))
                    print(f"[警告] レート制限 ({parent['thread_ts']}): "
                          f"{wait}秒待機して再試行", file=sys.stderr)
                    time.sleep(wait + 1)
                    try:
                        replies = client.conversations_replies(
                            channel=channel_id, ts=parent["thread_ts"], limit=200
                        )
                        for reply in replies["messages"][1:]:
                            reply["_parent_msg"] = parent
                            extra.append(reply)
                        continue
                    except SlackApiError as e2:
                        err = e2.response.get("error", "")
                url = thread_url(parent)
                print(f"[警告] スレッド取得失敗 ({err}): {url}", file=sys.stderr)
            if (i + 1) % 100 == 0:
                time.sleep(1)
        messages.extend(extra)
    return messages


def extract_dates(text: str, default_year: int) -> set[date]:
    """テキストから日付を抽出。月日のみ(年なし)の表記は default_year を補完する。
    対応書式:
      - 2026-05-20 / 2026/05/20 / 2026.05.20
      - 2026年5月20日
      - 5月20日 (年なし → default_year)
      - 5/20    (年なし → default_year)
    """
    found: set[date] = set()
    if not text:
        return found

    # YYYY-MM-DD / YYYY/MM/DD / YYYY.MM.DD / YYYY年M月D日
    for m in re.finditer(r'(?<!\d)(\d{4})[-/.年](\d{1,2})[-/.月](\d{1,2})日?(?!\d)', text):
        try:
            found.add(date(int(m.group(1)), int(m.group(2)), int(m.group(3))))
        except ValueError:
            pass

    # YYYYMMDD (区切りなし8桁)
    for m in re.finditer(r'(?<!\d)(\d{4})(\d{2})(\d{2})(?!\d)', text):
        try:
            found.add(date(int(m.group(1)), int(m.group(2)), int(m.group(3))))
        except ValueError:
            pass

    # M月D日 (年なし)
    for m in re.finditer(r'(?<!\d)(\d{1,2})月(\d{1,2})日', text):
        try:
            found.add(date(default_year, int(m.group(1)), int(m.group(2))))
        except ValueError:
            pass

    # M/D (年なし、前後が数字でないとき)
    for m in re.finditer(r'(?<!\d)(\d{1,2})/(\d{1,2})(?!\d)', text):
        mm, dd = int(m.group(1)), int(m.group(2))
        if 1 <= mm <= 12 and 1 <= dd <= 31:
            try:
                found.add(date(default_year, mm, dd))
            except ValueError:
                pass

    return found


NON_TARGET_KNOWN_EXTS = {
    ".pdf", ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".svg",
    ".pptx", ".ppt", ".txt", ".zip", ".tar", ".gz",
    ".mp4", ".mov", ".avi", ".mp3", ".wav",
}


def collect_teams_posts(
    messages: list[dict],
    target_date: "date | None" = None,
) -> list[dict]:
    """親メッセージのうち Teams URL を含むものだけ抽出。
    target_date が指定された場合、Slack投稿日が ±14日以内のものに絞り込む
    (実際の日付フィルタは Teams会議の実日付で行うため、ここは粗いフィルタ)。
    """
    from datetime import timezone, timedelta
    JST = timezone(timedelta(hours=9))
    results: list[dict] = []
    seen_ts: set[str] = set()

    for msg in messages:
        parent = msg.get("_parent_msg") or msg
        ts = parent.get("ts", "")
        if not ts or ts in seen_ts:
            continue
        if not extract_teams_urls(parent):
            continue
        if target_date:
            try:
                posted_jst = datetime.fromtimestamp(float(ts), tz=JST)
                days_diff = abs((posted_jst.date() - target_date).days)
                # Teams会議の実日付で最終判定するので、ここでは±14日まで許容
                if days_diff > 14:
                    continue
            except (ValueError, TypeError):
                continue
        seen_ts.add(ts)
        snippet = (parent.get("text") or "").splitlines()[0][:40] \
            if parent.get("text") else ts
        results.append({
            "_msg": parent,
            "_posted_ts": ts,
            "name": f"[teams投稿] {snippet}",
        })
    return results


def collect_files(messages: list[dict], client: WebClient,
                  target_date: date | None = None,
                  name_pattern: "re.Pattern | None" = None) -> list[dict]:
    """メッセージ群から対象ファイルを抜き出す。
    msg.files の添付、Slackファイルリンク、SharePoint URLを統合して扱う。
    SharePoint URLは末尾が実ファイル名でないことがあるので、明らかに非対象な
    拡張子のものだけ除外し、それ以外はGraph APIで実名を確認するために通過させる。
    """
    found = []
    seen_ids: set[str] = set()
    default_year = target_date.year if target_date else datetime.now().year

    for msg in messages:
        text = msg.get("text", "") or ""
        msg_dates = extract_dates(text, default_year) if target_date else set()

        attached = list(msg.get("files", []) or [])

        for fid in extract_referenced_file_ids(msg):
            if fid in seen_ids:
                continue
            try:
                resp = client.files_info(file=fid)
                attached.append(resp["file"])
            except SlackApiError as e:
                print(f"[警告] files.info 失敗 (id={fid}): "
                      f"{e.response.get('error')}", file=sys.stderr)

        for sp_url in extract_sharepoint_urls(msg):
            tentative_name = filename_from_sharepoint_url(sp_url)
            sp_entry = {
                "name": tentative_name,
                "user": msg.get("user", ""),
                "_sharepoint_url": sp_url,
                "_msg_dates": msg_dates,
            }
            attached.append(sp_entry)

        for f in attached:
            fid = f.get("id", "")
            if fid and fid in seen_ids:
                continue
            name = f.get("name") or ""
            ext = Path(name).suffix.lower()
            is_sharepoint = bool(f.get("_sharepoint_url"))

            # 拡張子チェック
            if is_sharepoint:
                # SharePointは実ファイル名がGraph問い合わせまで分からないので、
                # 明らかに非対象な拡張子だけ除外
                if ext in NON_TARGET_KNOWN_EXTS:
                    continue
            else:
                if ext not in TARGET_EXTS:
                    continue
                # ファイル名パターンチェック (SlackファイルはここでOK)
                if name_pattern and not name_pattern.search(name):
                    continue

            # 日付フィルタ
            if target_date is not None:
                if is_sharepoint:
                    # SharePointは本文の日付のみで判定 (実ファイル名側は後で再判定)
                    if target_date not in msg_dates:
                        continue
                else:
                    file_dates = extract_dates(name, default_year)
                    if target_date not in (msg_dates | file_dates):
                        continue

            if fid:
                seen_ids.add(fid)
            f["_posted_ts"] = msg.get("ts", "")
            f["_msg"] = msg
            found.append(f)
    return found


def download_file(file_obj: dict, token: str) -> bytes:
    url = file_obj.get("url_private_download") or file_obj["url_private"]
    headers = {"Authorization": f"Bearer {token}"}
    r = requests.get(url, headers=headers, timeout=60)
    r.raise_for_status()
    return r.content


# ---------- Microsoft Graph / SharePoint ----------

def get_graph_token(extra_scopes: list[str] | None = None) -> str:
    """Graph APIアクセストークンを取得。
    キャッシュがあれば自動更新、無ければデバイスコードフロー認証。
    extra_scopes で追加スコープを要求可能 (OnlineMeetings.Read 等)。
    """
    import msal
    cache = msal.SerializableTokenCache()
    if os.path.exists(MSAL_CACHE_FILE):
        with open(MSAL_CACHE_FILE) as f:
            cache.deserialize(f.read())

    app = msal.PublicClientApplication(
        GRAPH_CLIENT_ID,
        authority="https://login.microsoftonline.com/organizations",
        token_cache=cache,
    )
    scopes = ["Files.Read.All"] + list(extra_scopes or [])

    result = None
    accounts = app.get_accounts()
    if accounts:
        result = app.acquire_token_silent(scopes, account=accounts[0])

    if not result:
        flow = app.initiate_device_flow(scopes=scopes)
        if "user_code" not in flow:
            raise RuntimeError(f"Device flow開始失敗: {flow}")
        print("\n=== Microsoft認証が必要です ===", file=sys.stderr)
        print(flow["message"], file=sys.stderr)
        print(f"要求スコープ: {scopes}", file=sys.stderr)
        print("===============================\n", file=sys.stderr)
        result = app.acquire_token_by_device_flow(flow)

    if "access_token" not in result:
        raise RuntimeError(
            f"トークン取得失敗: {result.get('error_description', result)}"
        )

    if cache.has_state_changed:
        with open(MSAL_CACHE_FILE, "w") as f:
            f.write(cache.serialize())

    return result["access_token"]


def download_sharepoint_file(url: str, graph_token: str) -> tuple[str, bytes]:
    """SharePoint共有URLからファイルをDL。(filename, content_bytes) を返す。"""
    import base64
    encoded = base64.urlsafe_b64encode(url.encode("utf-8")).decode("ascii").rstrip("=")
    share_id = "u!" + encoded
    headers = {"Authorization": f"Bearer {graph_token}"}

    # メタ情報 (正式なファイル名取得用)
    filename = ""
    try:
        meta = requests.get(
            f"https://graph.microsoft.com/v1.0/shares/{share_id}/driveItem",
            headers=headers, timeout=30,
        )
        if meta.ok:
            filename = meta.json().get("name", "") or ""
    except requests.RequestException:
        pass

    # 本体ダウンロード
    resp = requests.get(
        f"https://graph.microsoft.com/v1.0/shares/{share_id}/driveItem/content",
        headers=headers, timeout=120, allow_redirects=True,
    )
    resp.raise_for_status()
    return filename, resp.content


def load_logs_json(path: str) -> tuple[list, "date | None"]:
    """logs.jsonを読み込んで (entries, max_posted_date) を返す。
    壊れていれば ([], None) を返す。
    """
    if not os.path.exists(path):
        return [], None
    try:
        with open(path, "rb") as f:
            raw = f.read()
        # BOM対応
        if raw.startswith(b'\xff\xfe'):
            content = raw.decode("utf-16-le", errors="replace").lstrip("\ufeff").strip()
        elif raw.startswith(b'\xfe\xff'):
            content = raw.decode("utf-16-be", errors="replace").lstrip("\ufeff").strip()
        elif raw.startswith(b'\xef\xbb\xbf'):
            content = raw[3:].decode("utf-8", errors="replace").strip()
        else:
            content = raw.decode("utf-8", errors="replace").strip()
        if not content:
            return [], None
        entries = json.loads(content)
        if not isinstance(entries, list):
            return [], None

        max_date: "date | None" = None
        for e in entries:
            if not isinstance(e, dict):
                continue
            posted = e.get("posted_at") or ""
            if not posted:
                continue
            try:
                # ISO形式想定 (e.g. "2026-05-29T14:00:00+09:00")
                dt = datetime.fromisoformat(posted.replace("Z", "+00:00"))
                d = dt.date()
                if max_date is None or d > max_date:
                    max_date = d
            except ValueError:
                continue
        return entries, max_date
    except Exception as e:
        print(f"[警告] logs.json読み込み失敗: {e}", file=sys.stderr)
        return [], None


def fetch_track_calendars(
    driver_emails: list[str],
    graph_token: str,
    target_date: "date | None",
    days_window: int = 14,
) -> dict:
    """各ドライバーのカレンダーから±N日分のオンライン会議を取得し、
    JoinUrl をキーにした dict を返す: {joinUrl: event_data}
    """
    from datetime import timedelta as _td
    headers = {"Authorization": f"Bearer {graph_token}"}
    if target_date is None:
        target_date = datetime.now().date()
    start = (datetime.combine(target_date, datetime.min.time())
             - _td(days=days_window)).isoformat() + "Z"
    end = (datetime.combine(target_date, datetime.min.time())
           + _td(days=days_window)).isoformat() + "Z"

    join_url_map: dict = {}
    for driver in driver_emails:
        try:
            r = requests.get(
                f"https://graph.microsoft.com/v1.0/users/{driver}/calendarView"
                f"?startDateTime={start}&endDateTime={end}"
                f"&$select=subject,start,end,onlineMeeting,isOnlineMeeting,organizer"
                f"&$top=500",
                headers=headers, timeout=60,
            )
            if not r.ok:
                print(f"[track-cal] {driver} 取得失敗 status={r.status_code}: "
                      f"{r.text[:200]}", file=sys.stderr)
                continue
            events = r.json().get("value", [])
            online_count = 0
            for ev in events:
                if not ev.get("isOnlineMeeting"):
                    continue
                join = (ev.get("onlineMeeting") or {}).get("joinUrl", "")
                if join and join not in join_url_map:
                    join_url_map[join] = ev
                    online_count += 1
            print(f"[track-cal] {driver}: 全{len(events)}件 / "
                  f"オンライン会議 {online_count}件を取得", file=sys.stderr)
        except requests.RequestException as e:
            print(f"[track-cal] {driver} 例外: {type(e).__name__}: {e}",
                  file=sys.stderr)
    return join_url_map


def get_teams_meeting(url: str, graph_token: str) -> dict | None:
    """Teams会議URLから会議情報を取得。
    1. OnlineMeetings API (自分主催の会議のみ)
    2. CalendarView API (自分のカレンダーにある予定)
    の順で試行。
    """
    headers = {"Authorization": f"Bearer {graph_token}"}
    safe_url = url.replace("'", "''")

    # === 1. OnlineMeetings API (自分主催) ===
    try:
        r = requests.get(
            f"https://graph.microsoft.com/v1.0/me/onlineMeetings"
            f"?$filter=JoinWebUrl eq '{safe_url}'",
            headers=headers, timeout=30,
        )
        if r.ok:
            items = r.json().get("value", [])
            if items:
                return items[0]
            # 0件 → カレンダーへフォールバック
        elif r.status_code == 403:
            # 主催者でない → カレンダーへフォールバック
            pass
        else:
            print(f"    [teams API失敗] status={r.status_code} url={url[:80]}",
                  file=sys.stderr)
            print(f"    [teams API失敗] body={r.text[:300]}", file=sys.stderr)
    except requests.RequestException as e:
        print(f"    [teams API例外] {type(e).__name__}: {e}", file=sys.stderr)

    # === 2. CalendarView API フォールバック (自分のカレンダー上の予定) ===
    from datetime import timedelta as _td
    now = datetime.utcnow()
    cal_start = (now - _td(days=90)).isoformat() + "Z"
    cal_end = (now + _td(days=90)).isoformat() + "Z"
    try:
        r = requests.get(
            f"https://graph.microsoft.com/v1.0/me/calendarView"
            f"?startDateTime={cal_start}&endDateTime={cal_end}"
            f"&$select=subject,start,end,onlineMeeting,isOnlineMeeting"
            f"&$top=500",
            headers=headers, timeout=60,
        )
        if not r.ok:
            print(f"    [calendar API失敗] status={r.status_code} "
                  f"body={r.text[:200]}", file=sys.stderr)
            return None
        events = r.json().get("value", [])
        for ev in events:
            if not ev.get("isOnlineMeeting"):
                continue
            join = (ev.get("onlineMeeting") or {}).get("joinUrl", "")
            if not join:
                continue
            if join == url or url in join or join in url:
                # CalendarViewの時刻形式をOnlineMeetingsの形式に揃える
                s_dt_str = ev["start"]["dateTime"]
                e_dt_str = ev["end"]["dateTime"]
                # 末尾のマイクロ秒を切り捨て、Zを付ける
                s_dt_str = s_dt_str.split(".")[0] + "Z" if "." in s_dt_str \
                    else (s_dt_str if s_dt_str.endswith("Z") else s_dt_str + "Z")
                e_dt_str = e_dt_str.split(".")[0] + "Z" if "." in e_dt_str \
                    else (e_dt_str if e_dt_str.endswith("Z") else e_dt_str + "Z")
                print(f"    [calendar] マッチ: {ev.get('subject', '')[:60]}",
                      file=sys.stderr)
                return {
                    "subject": ev.get("subject", ""),
                    "startDateTime": s_dt_str,
                    "endDateTime": e_dt_str,
                    "_source": "calendarView",
                }
        print(f"    [calendar] joinUrl一致せず ({len(events)}件中): {url[:80]}",
              file=sys.stderr)
        return None
    except requests.RequestException as e:
        print(f"    [calendar API例外] {type(e).__name__}: {e}", file=sys.stderr)
        return None


def parse_teams_subject(subject: str) -> dict:
    """Teams会議件名から Trackname / Customer / Route を抽出。
    形式: 【GIGA■■】○○様 実証 往路 関東→関西
    """
    result: dict = {}
    if not subject:
        return result

    # Trackname & Track-num: 【...】
    m = TRACKNAME_RE.search(subject)
    if m:
        track = m.group(1).strip()
        result["Trackname"] = track
        nm = TRACKNUM_RE.search(track)
        if nm:
            result["Track-num"] = nm.group(1)

    # 】の後を Customer + Route に分割
    m = re.search(r'】(.+)$', subject)
    if m:
        tail = m.group(1).strip()
        tokens = tail.split()
        # 末尾から → を含むトークンを Route として分離
        route_tokens = []
        while tokens and CUSTOMER_TAIL_ARROW_RE.search(tokens[-1]):
            route_tokens.insert(0, tokens.pop())
        # 末尾装飾を除去
        while tokens and re.fullmatch(r'[※#\-=●○◯◎★☆]+', tokens[-1]):
            tokens.pop()
        result["Customer"] = " ".join(tokens)
        result["Route"] = " ".join(route_tokens)
    return result


# ---------- Extractors ----------

def extract_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    out = []
    for para in doc.paragraphs:
        if para.text.strip():
            out.append(para.text)
    for table in doc.tables:
        out.append("")  # 区切り
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            out.append(" | ".join(cells))
    return "\n".join(out)


def extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    out = []
    for sheet in wb.worksheets:
        out.append(f"--- Sheet: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            if any(c is not None and str(c).strip() for c in row):
                out.append("\t".join("" if c is None else str(c) for c in row))
    return "\n".join(out)


def extract_xls(data: bytes) -> str:
    import pandas as pd
    sheets = pd.read_excel(io.BytesIO(data), sheet_name=None, header=None, engine="xlrd")
    out = []
    for name, df in sheets.items():
        out.append(f"--- Sheet: {name} ---")
        out.append(df.to_csv(sep="\t", index=False, header=False).rstrip())
    return "\n".join(out)


def extract_csv(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp932"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


EXTRACTORS = {
    ".docx": extract_docx,
    ".xlsx": extract_xlsx,
    ".xls": extract_xls,
    ".csv": extract_csv,
}


# ---------- Structured extractors (運行記録など) ----------

TIME_RE = re.compile(r'^\s*(\d{1,2}):(\d{2})\s*$')
DRIVING_DATE_RE = re.compile(r'(\d{4})[/\-.年](\d{1,2})[/\-.月](\d{1,2})')
WEEKDAY_CHARS = {"月", "火", "水", "木", "金", "土", "日"}


def parse_driving_log(data: bytes) -> dict:
    """1ファイルに複数日が横並びで入っている運行記録を解析。

    各日ブロックは以下を持つ前提:
      - ヘッダ部: 日付 / 曜日 / ルート / ADルート設定 / 燃料
      - 「時間」「業務」「メモ」の3列ヘッダ
      - その下に時系列の運行データ

    返り値:
        {"days": [
            {date, day_of_week, route, ad_route, fuel,
             entries:[{time,task,memo},...],
             ad_start?, ad_end?, breaks?}
        ]}
    """
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=False)
    days = []

    for sheet in wb.worksheets:
        # 全セルを (row, col) -> value にマップ
        cells = {}
        max_row = 0
        for r_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            for c_idx, val in enumerate(row, start=1):
                if val is not None:
                    cells[(r_idx, c_idx)] = val
                    max_row = max(max_row, r_idx)

        # "時間"ヘッダ + 同じ行の "業務" "メモ" を見つける
        time_headers = []
        for (r, c), v in cells.items():
            if not (isinstance(v, str) and v.strip() == "時間"):
                continue
            task_col = memo_col = None
            for offset in range(1, 5):
                neighbor = cells.get((r, c + offset))
                if isinstance(neighbor, str):
                    if "業務" in neighbor and task_col is None:
                        task_col = c + offset
                    if "メモ" in neighbor and memo_col is None:
                        memo_col = c + offset
            if task_col and memo_col:
                time_headers.append((r, c, task_col, memo_col))

        for header_row, time_col, task_col, memo_col in time_headers:
            # ブロック上方を最大10行スキャンして日付・ルート情報を取得
            info = {"date": None, "day_of_week": None,
                    "route": None, "ad_route": None, "fuel": None}
            for r in range(header_row - 1, max(0, header_row - 11), -1):
                for c in range(time_col, memo_col + 1):
                    v = cells.get((r, c))
                    if v is None:
                        continue
                    if isinstance(v, str):
                        s = v.strip()
                        m = DRIVING_DATE_RE.search(s)
                        if m and not info["date"]:
                            info["date"] = (
                                f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"
                            )
                        if s in WEEKDAY_CHARS and not info["day_of_week"]:
                            info["day_of_week"] = s
                        if ("支線" in s or "幹線" in s) and not info["route"]:
                            info["route"] = s
                        if "ADルート" in s and not info["ad_route"]:
                            info["ad_route"] = s
                        if s.startswith("燃料") and not info["fuel"]:
                            info["fuel"] = s.split(":", 1)[-1].split(":", 1)[-1].strip()
                    elif hasattr(v, "year") and hasattr(v, "month"):
                        if not info["date"]:
                            info["date"] = v.strftime("%Y-%m-%d")

            # データ行
            entries = []
            ad_start = ad_end = None
            breaks = []
            for r in range(header_row + 1, max_row + 1):
                tv = cells.get((r, time_col))
                if tv is None:
                    continue
                if hasattr(tv, "hour") and hasattr(tv, "minute"):
                    time_str = f"{tv.hour}:{tv.minute:02d}"
                else:
                    m = TIME_RE.match(str(tv))
                    if not m:
                        continue
                    time_str = f"{int(m.group(1))}:{int(m.group(2)):02d}"

                task_v = cells.get((r, task_col))
                memo_v = cells.get((r, memo_col))
                task = str(task_v).strip() if task_v is not None else ""
                memo = str(memo_v).strip() if memo_v is not None else ""

                entries.append({"time": time_str, "task": task, "memo": memo})

                if "AD開始" in memo:
                    ad_start = {"time": time_str, "location": task}
                elif "AD終了" in memo:
                    ad_end = {"time": time_str, "location": task}
                elif "休憩" in memo:
                    breaks.append({"time": time_str, "location": task, "note": memo})

            if entries:
                day_obj = {**info, "entries": entries}
                if ad_start:
                    day_obj["ad_start"] = ad_start
                if ad_end:
                    day_obj["ad_end"] = ad_end
                if breaks:
                    day_obj["breaks"] = breaks
                days.append(day_obj)

    return {"days": days}


STRUCTURED_EXTRACTORS = {
    "driving": (parse_driving_log, {".xlsx"}),
}


# ---------- Tracking metadata extractors (Slackスレッド本文 → tracking) ----------

# 半角/全角の括弧・コロンを両方マッチさせる
_PAREN_O = r'[\(\uff08]'      # ( or (
_PAREN_C = r'[\)\uff09]'      # ) or )
_COLON   = r'[:\uff1a]'        # : or :

TRACKING_DATE_TITLE_RE = re.compile(r'(\d{1,2})月(\d{1,2})日')
TRACKNAME_RE = re.compile(r'【([^】]+)】')
TRACKNUM_RE = re.compile(r'(\d+)')
CUSTOMER_TAIL_ARROW_RE = re.compile(r'→')
DRIVER_LINE_RE = re.compile(
    rf'ドライバー\s*{_PAREN_O}幹線{_PAREN_C}\s*{_COLON}\s*<@(U[A-Z0-9]+)>'
)
OPERATOR_LINE_RE = re.compile(
    rf'オペレータ\s*{_PAREN_O}幹線{_PAREN_C}\s*{_COLON}\s*<@(U[A-Z0-9]+)>'
)
SW_VER_RE = re.compile(
    rf'SW\s*Ver\.?\s*{_COLON}\s*WIP\s*{_COLON}\s*(v\S+)', re.IGNORECASE
)
ROUTE_LINE_RE = re.compile(rf'自動運転区間\s*{_COLON}\s*([^\n]+)')


def resolve_user_name(user_id: str, client: WebClient, cache: dict) -> str:
    """ユーザーIDから表示名を取得。
    優先順: profile.display_name → profile.real_name → real_name → name → <@U...>
    """
    if user_id in cache:
        return cache[user_id]
    try:
        resp = client.users_info(user=user_id)
        user = resp["user"]
        prof = user.get("profile", {}) or {}
        # display_name は profile.display_name が正(ユーザーが設定した表示名)
        name = (prof.get("display_name") or "").strip() \
            or (prof.get("real_name") or "").strip() \
            or (user.get("real_name") or "").strip() \
            or user.get("name") \
            or f"<@{user_id}>"
        cache[user_id] = name
        return name
    except SlackApiError as e:
        err = e.response.get("error", "unknown")
        # 初回失敗時は warning を出す (同じユーザーで何度も出ないようキャッシュに記録)
        print(f"[警告] ユーザー解決失敗 ({user_id}): {err}"
              + (" -- users:read スコープを付与して再インストール必要"
                 if err == "missing_scope" else ""),
              file=sys.stderr)
        cache[user_id] = f"<@{user_id}>"
        return cache[user_id]


def normalize_slack_text(text: str) -> str:
    """Slack特有の装飾を除去 (絵文字ショートコード、太字、リンク表記など)"""
    if not text:
        return ""
    # <url|display> → display, <url> → url
    text = re.sub(r'<([^|>\s]+)\|([^>]+)>', r'\2', text)
    text = re.sub(r'<(https?://[^>]+)>', r'\1', text)
    # メンション <@Uxxx> は残す
    # 絵文字 :emoji_name: を削除
    #   前が英数字/ピリオド(=ラベル系)でない場合は絵文字として除去
    #   日本語文字の後の絵文字も除去できる
    text = re.sub(r'(?<![a-zA-Z0-9_.]):[a-zA-Z0-9_+\-]+:(?!//)', '', text)
    # *bold* / _italic_ / ~strike~
    text = re.sub(r'[\*_~]+', '', text)
    return text


def extract_tracking_metadata(msg: dict, client: WebClient, user_cache: dict,
                              default_year: int) -> dict:
    """親メッセージの本文から tracking 用フィールドを抽出。
    msg がスレッド返信なら、_parent_msg にあたる親メッセージの本文を見る。"""
    source_msg = msg.get("_parent_msg") or msg
    raw_text = source_msg.get("text", "") or ""
    text = normalize_slack_text(raw_text)
    result = {}

    # タイトル行 = 【...】を含む最初の行 (見出し行の後ろにあるケース対応)
    title_line = ""
    for line in text.split("\n"):
        if "【" in line and "】" in line:
            title_line = line
            break
    if not title_line:
        title_line = text.split("\n", 1)[0]
    first_line = title_line

    # Trackname & Track-num: 【...】
    m = TRACKNAME_RE.search(first_line)
    if m:
        track = m.group(1).strip()
        result["Trackname"] = track
        nm = TRACKNUM_RE.search(track)
        if nm:
            result["Track-num"] = nm.group(1)

    # Customer: 】 の後 ～ 末尾の "X→Y" (方向)を除いた部分
    # 装飾は normalize で除去済み。※ などのセクション区切りで停止
    m = re.search(r'】([^\n<※]+)', first_line)
    if m:
        tail = m.group(1).strip()
        tokens = tail.split()
        # 末尾から、方向(→を含む)のトークンを取り除く
        while tokens and CUSTOMER_TAIL_ARROW_RE.search(tokens[-1]):
            tokens.pop()
        # 末尾の装飾も削除
        while tokens and re.fullmatch(r'[※#\-=●○◯◎★☆]+', tokens[-1]):
            tokens.pop()
        result["Customer"] = " ".join(tokens)

    # タイトル内の日付 (YY/MM/DD)
    m = TRACKING_DATE_TITLE_RE.search(first_line)
    if m:
        month, day = int(m.group(1)), int(m.group(2))
        try:
            result["_title_date"] = date(default_year, month, day)
        except ValueError:
            pass

    # Driver-name: メンション解決
    m = DRIVER_LINE_RE.search(text)
    if m:
        result["Driver-name"] = resolve_user_name(m.group(1), client, user_cache)

    # Operator: 記載なし(=行ごと無い)場合は "oneman"
    m = OPERATOR_LINE_RE.search(text)
    if m:
        result["Operator"] = resolve_user_name(m.group(1), client, user_cache)
    else:
        result["Operator"] = "oneman"

    # SW-ver
    m = SW_VER_RE.search(text)
    if m:
        result["SW-ver"] = m.group(1)

    # Route: 自動運転区間：...(行末まで)
    m = ROUTE_LINE_RE.search(text)
    if m:
        result["Route"] = m.group(1).strip()

    return result


def get_message_permalink(client: WebClient, channel_id: str, ts: str) -> str:
    try:
        resp = client.chat_getPermalink(channel=channel_id, message_ts=ts)
        return resp.get("permalink", "")
    except SlackApiError:
        return ""


def build_labeled_record(meta: dict, day_data: dict | None, url: str) -> dict:
    """meta + Excel該当日 → ラベル付きの dict 形式を作る (人間が読みやすい版)"""
    title_date: date | None = meta.get("_title_date")
    date_str = title_date.strftime("%Y/%m/%d") if title_date else ""

    teams_tr = meta.get("_teams_time_range")
    if teams_tr:
        time_range = teams_tr
    else:
        start_time = end_time = ""
        if day_data:
            s = day_data.get("ad_start") or {}
            e = day_data.get("ad_end") or {}
            start_time = s.get("time", "") if s else ""
            end_time = e.get("time", "") if e else ""
        time_range = ""
        if date_str and start_time and end_time:
            time_range = f"{date_str}T{start_time}+9:00/{date_str}T{end_time}+9:00"

    track_num = meta.get("Track-num", "")
    return {
        "Trackname": meta.get("Trackname", ""),
        "Track-num|YY/MM/DD": (f"{track_num}|{date_str}" if date_str else track_num),
        "Time-range": time_range,
        "Driver": meta.get("Driver-name", ""),
        "operator": meta.get("Operator", "oneman"),
        "SW-version": meta.get("SW-ver", ""),
        "selfdrive section": meta.get("Route", ""),
        "loaded liggage": meta.get("Customer", ""),
        "url": url,
    }


def build_tracking_record(meta: dict, day_data: dict | None, posted_iso: str,
                          url: str) -> list:
    """meta + Excel該当日 → 目標フォーマット([str, str, str, dict]) を作る"""
    track = meta.get("Trackname", "")
    track_num = meta.get("Track-num", "")
    title_date: date | None = meta.get("_title_date")
    date_str = title_date.strftime("%Y/%m/%d") if title_date else ""

    # Teams由来のTime-rangeがあれば優先、無ければExcelから組み立て
    teams_tr = meta.get("_teams_time_range")
    if teams_tr:
        time_range = teams_tr
    else:
        start_time = end_time = ""
        if day_data:
            s = day_data.get("ad_start") or {}
            e = day_data.get("ad_end") or {}
            start_time = s.get("time", "") if s else ""
            end_time = e.get("time", "") if e else ""
        time_range = ""
        if date_str and start_time and end_time:
            time_range = f"{date_str}T{start_time}+9:00/{date_str}T{end_time}+9:00"

    return [
        track,
        f"{track_num}|{date_str}" if date_str else track_num,
        time_range,
        {
            "Driver": meta.get("Driver-name", ""),
            "operator": meta.get("Operator", "oneman"),
            "SW-version": meta.get("SW-ver", ""),
            "selfdrive section": meta.get("Route", ""),
            "loaded liggage": meta.get("Customer", ""),
            "url": url,
        },
    ]


# ---------- CLI ----------

def parse_since(s: str) -> str:
    """YYYY-MM-DD を unix秒(文字列)に変換"""
    return str(datetime.strptime(s, "%Y-%m-%d").timestamp())


def resolve_relative_date(value: str | None) -> "date | None":
    """日付文字列を date に変換。
    "today"/"yesterday"/"tomorrow"/"N_days_ago"/"N_days_later" の相対表現にも対応。
    """
    if not value:
        return None
    from datetime import timedelta
    today = datetime.now().date()
    if value == "today":
        return today
    if value == "yesterday":
        return today - timedelta(days=1)
    if value == "tomorrow":
        return today + timedelta(days=1)
    m = re.match(r'(\d+)_days?_ago$', value)
    if m:
        return today - timedelta(days=int(m.group(1)))
    m = re.match(r'(\d+)_days?_later$', value)
    if m:
        return today + timedelta(days=int(m.group(1)))
    return datetime.strptime(value, "%Y-%m-%d").date()


def build_legs_record(meta: dict, day_data: dict | None, url: str) -> list:
    """legs.json用レコード (配列形式)。
    形式: [Trackname, "Track-num|YY/MM/DD",
           "YYYY-MM-DDTHH:MM:SS.000+09:00/YYYY-MM-DDTHH:MM:SS.000+09:00",
           {SW-version, selfdrive_section, loaded_luggage, url}]
    Time-range は Teams会議の実日付を使う(跨日OK)。
    """
    from datetime import timezone, timedelta, time as dtime
    JST = timezone(timedelta(hours=9))

    track = meta.get("Trackname", "")
    track_num = meta.get("Track-num", "")
    title_date: "date | None" = meta.get("_title_date")
    date_str = title_date.strftime("%Y/%m/%d") if title_date else ""

    teams_start = meta.get("_teams_start_dt")
    teams_end = meta.get("_teams_end_dt")

    time_range = ""
    if teams_start and teams_end:
        # Teams の実日付・実時刻でフルISO出力 (跨日対応)
        time_range = (
            f"{teams_start.isoformat(timespec='milliseconds')}/"
            f"{teams_end.isoformat(timespec='milliseconds')}"
        )
    elif title_date and day_data:
        # Excelフォールバック
        s_time = (day_data.get("ad_start") or {}).get("time", "")
        e_time = (day_data.get("ad_end") or {}).get("time", "")
        if s_time and e_time:
            try:
                sh, sm = (int(x) for x in s_time.split(":")[:2])
                eh, em = (int(x) for x in e_time.split(":")[:2])
                s_dt = datetime.combine(title_date, dtime(sh, sm), tzinfo=JST)
                e_dt = datetime.combine(title_date, dtime(eh, em), tzinfo=JST)
                # 終了 < 開始 なら跨日扱い
                if e_dt < s_dt:
                    e_dt = e_dt + timedelta(days=1)
                time_range = (
                    f"{s_dt.isoformat(timespec='milliseconds')}/"
                    f"{e_dt.isoformat(timespec='milliseconds')}"
                )
            except (ValueError, IndexError):
                pass

    return [
        track,
        f"{track_num}|{date_str}" if date_str else track_num,
        time_range,
        {
            "SW-version": meta.get("SW-ver", ""),
            "selfdrive_section": meta.get("Route", ""),
            "loaded_luggage": meta.get("Customer", ""),
            "url": url,
        },
    ]


def is_legs_record_complete(rec) -> tuple[bool, list]:
    """legsレコードの全フィールドが埋まっているかチェック。
    戻り値: (完全か, 空フィールド名のリスト)
    "---" "--" "-" "未定" "TBD" "なし" 等のプレースホルダも「未取得」扱い。
    """
    # プレースホルダパターン (空白除去後にこれらにマッチしたら「未取得」)
    PLACEHOLDER_PATTERNS = {
        "", "-", "--", "---", "----", "−", "—", "ー",
        "未定", "未取得", "未設定", "なし", "無し", "tbd", "TBD",
        "n/a", "N/A", "na", "NA", "なし", "?", "？", "不明",
    }

    def is_placeholder(value) -> bool:
        if value is None:
            return True
        s = str(value).strip()
        if s in PLACEHOLDER_PATTERNS:
            return True
        # ハイフンだけで構成されているケース (---- など長さ可変)
        if s and all(ch in "-−—ー" for ch in s):
            return True
        return False

    missing: list = []
    if not isinstance(rec, list) or len(rec) < 4:
        return False, ["record_format"]

    if is_placeholder(rec[0]):
        missing.append("Trackname")
    if not rec[1] or "|" not in (rec[1] or ""):
        missing.append("Date")
    if is_placeholder(rec[2]):
        missing.append("Time-range")

    meta = rec[3] if isinstance(rec[3], dict) else {}
    for key in ["SW-version", "selfdrive_section", "loaded_luggage", "url"]:
        if is_placeholder(meta.get(key)):
            missing.append(key)

    return len(missing) == 0, missing


def legs_dedup_key(rec) -> tuple:
    """legs.json用の重複判定キー: (Trackname, 日付, 往路/復路)
    新旧両形式(配列 / dict) を受け付ける。
    """
    if isinstance(rec, list) and len(rec) >= 4:
        track = rec[0] or ""
        date_part = rec[1] or ""
        date_str = date_part.split("|", 1)[1] if "|" in date_part else ""
        d = rec[3] if isinstance(rec[3], dict) else {}
        liggage = d.get("loaded_luggage") or d.get("loaded liggage", "") or ""
    elif isinstance(rec, dict):
        track = rec.get("Trackname", "") or ""
        date_part = rec.get("Track-num|YY/MM/DD", "") or ""
        date_str = date_part.split("|", 1)[1] if "|" in date_part else ""
        liggage = rec.get("loaded_luggage") or rec.get("loaded liggage", "") or ""
    else:
        return ("", "", "")

    direction = ""
    for word in ("往路", "復路"):
        if word in liggage:
            direction = word
            break
    return (track, date_str, direction)


def send_slack_notification(webhook_url: str, text: str) -> None:
    """Slack Incoming Webhook にメッセージを送信"""
    try:
        r = requests.post(webhook_url, json={"text": text}, timeout=15)
        if not r.ok:
            print(f"[警告] Slack通知失敗 ({r.status_code}): {r.text[:200]}",
                  file=sys.stderr)
    except requests.RequestException as e:
        print(f"[警告] Slack通知送信エラー: {e}", file=sys.stderr)


def format_tracking_record_for_slack(rec_array: list) -> str:
    """tracking レコード([Trackname, ..., {meta}]) を通知用テキストに整形"""
    track = rec_array[0] if len(rec_array) > 0 else ""
    meta = rec_array[3] if len(rec_array) > 3 and isinstance(rec_array[3], dict) else {}
    customer = meta.get("loaded liggage", "")
    section = meta.get("selfdrive section", "")
    url = meta.get("url", "")

    lines = [f"*{track}*"] if track else []
    if customer:
        lines.append(f"  運行名: {customer}")
    if section:
        lines.append(f"  自動運転区間: {section}")
    if url:
        lines.append(f"  URL: <{url}|スレッドを開く>")
    return "\n".join(lines)


def main() -> None:
    # まず --config パスだけ先に取り出す (config.jsonがデフォルト)
    pre_parser = argparse.ArgumentParser(add_help=False)
    pre_parser.add_argument("--config", default="config.json")
    pre_args, _ = pre_parser.parse_known_args()

    config: dict = {}
    if os.path.exists(pre_args.config):
        try:
            with open(pre_args.config, encoding="utf-8") as f:
                config = __import__("json").load(f)
            print(f"[設定] {pre_args.config} を読み込み ({len(config)}項目)",
                  file=sys.stderr)
        except Exception as e:
            sys.exit(f"設定ファイル読み込み失敗 ({pre_args.config}): {e}")

    parser = argparse.ArgumentParser(
        description="Slackチャンネルに添付された Excel / Word / CSV を読み出す"
    )
    parser.add_argument("--config", default="config.json",
                        help="設定ファイル(JSON)のパス (default: config.json)")
    parser.add_argument("--channel", nargs="+",
                        help="チャンネルID (C0123...) もしくは名前 (#general)。"
                             "複数指定可 (例: --channel C01ABC C02DEF)。"
                             "config.json では文字列または配列で指定可")
    parser.add_argument("--limit", type=int, default=20,
                        help="取得する親メッセージ件数 (default: 20)")
    parser.add_argument("--since", help="この日付以降に投稿されたメッセージのみ (YYYY-MM-DD)")
    parser.add_argument("--content-date",
                        help="本文 or 添付ファイル名にこの日付が書かれた投稿のみ対象 (YYYY-MM-DD)")
    parser.add_argument("--threads", action="store_true",
                        help="スレッド返信内の添付も対象にする")
    parser.add_argument("--extract", choices=["raw", "driving", "tracking"], default="raw",
                        help="raw=テキストダンプ / driving=運行記録JSON / "
                             "tracking=Slackスレッド+Excelから集約した1行レコード")
    parser.add_argument("--tracking-source", choices=["excel", "teams"], default="excel",
                        help="(trackingモード) excel=Excelから時刻取得(従来) / "
                             "teams=Teams会議から件名・時刻取得")
    parser.add_argument("--track-calendars", nargs="+", default=None,
                        help="(trackingモード+teams) "
                             "他人のカレンダーから会議情報を引くドライバーのメール/UPN。"
                             "カレンダー共有が必要(Calendars.Read.Shared)")
    parser.add_argument("--name-pattern", default=None,
                        help="ファイル名がこの正規表現にマッチするものだけ対象 "
                             "(例: \"^運行\" で運行で始まる名前のみ)")
    parser.add_argument("--memo-filter", default=None,
                        help="(driving モード) メモ列がこの正規表現にマッチする行だけ"
                             "残す (例: \"AD\" / \"休憩\" / \"AD|休憩\")")
    parser.add_argument("--debug", action="store_true",
                        help="取得した各メッセージの中身(添付/URL/本文先頭)を表示")
    parser.add_argument("--append", action="store_true",
                        help="(trackingモード) 既存のresult.jsonに追記。"
                             "スレッドURLが重複するレコードはスキップ")
    parser.add_argument("--legs-out", default=None,
                        help="(trackingモード) ラベル付きレコード(__labeled_view__相当) "
                             "を別のJSONファイルに保存。"
                             "Trackname+日付+往路/復路 で重複をスキップ")
    parser.add_argument("--logs-out", default=None,
                        help="(trackingモード) legs.jsonに書き込んだ投稿の "
                             "投稿日履歴を保存。次回実行時、この最新日より前は "
                             "Slackから取得しない (高速化)")
    parser.add_argument("--notify-webhook-url", default=None,
                        help="完了時にこのSlack Incoming WebhookへPOSTして結果通知")
    parser.add_argument("--start-notify-webhook-url", default=None,
                        help="開始時にこのSlack Incoming WebhookへPOSTして開始通知 "
                             "(個人DM用Webhookに設定すると、開始通知だけDMで受け取れる)")
    parser.add_argument("--out", help="結果をファイルに出力 (省略時は標準出力)")

    # configの未知キーを警告 (タイポ防止)
    known_dests = {a.dest for a in parser._actions}
    for key in config:
        if key not in known_dests:
            print(f"[警告] 設定ファイルの未知のキー: {key}", file=sys.stderr)

    # CLI > config の優先順位で merge
    parser.set_defaults(**{k: v for k, v in config.items() if k in known_dests})
    args = parser.parse_args()

    # === 診断: 実際に使われている主要な設定値 ===
    print(f"[診断] extract={args.extract!r}, tracking_source={args.tracking_source!r}, "
          f"content_date={args.content_date!r}, append={args.append}, "
          f"out={args.out!r}, legs_out={args.legs_out!r}",
          file=sys.stderr)

    if not args.channel:
        sys.exit("エラー: --channel もしくは config.json の \"channel\" 指定が必要です")

    # channel は単数(str) or 複数(list) を受け取る
    channels = args.channel if isinstance(args.channel, list) else [args.channel]
    channels = [c for c in channels if c]  # 空文字を除外

    token = get_token()
    client = WebClient(token=token)

    # レート制限エラー時に自動リトライ (retry-afterヘッダを尊重して待機)
    from slack_sdk.http_retry.builtin_handlers import RateLimitErrorRetryHandler
    client.retry_handlers.append(RateLimitErrorRetryHandler(max_retry_count=5))

    since_ts = parse_since(args.since) if args.since else None
    target_date = resolve_relative_date(args.content_date)
    if args.content_date and target_date:
        print(f"[情報] content_date = {target_date.isoformat()} "
              f"(指定: {args.content_date!r})", file=sys.stderr)

    # 開始通知 (個人DM等のWebhookに送信)
    if args.start_notify_webhook_url:
        start_lines = ["🚀 zp_summary の実行を開始しました"]
        if target_date:
            start_lines.append(f"対象日: {target_date.isoformat()}")
        start_lines.append(f"実行開始: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        if args.extract == "tracking":
            mode = args.tracking_source
            start_lines.append(f"モード: tracking ({mode})")
        send_slack_notification(args.start_notify_webhook_url, "\n".join(start_lines))

    # logs.json を読み込み、最新投稿日より新しい投稿のみ取得するよう since_ts を調整
    existing_logs: list = []
    log_max_date: "date | None" = None
    if args.logs_out:
        existing_logs, log_max_date = load_logs_json(args.logs_out)
        if log_max_date:
            log_ts = datetime.combine(log_max_date, datetime.min.time()).timestamp()
            if since_ts is None or float(since_ts) < log_ts:
                since_ts = str(log_ts)
                print(f"[logs] logs.json の最新投稿日 {log_max_date} を since として使用 "
                      f"(既存 {len(existing_logs)} 件)", file=sys.stderr)
            else:
                print(f"[logs] 既存 {len(existing_logs)} 件 / "
                      f"最新投稿日 {log_max_date} (--since の方が新しいため未使用)",
                      file=sys.stderr)
    def compile_filter(value):
        """文字列 or 文字列リスト を正規表現にコンパイル。リストはORで結合。"""
        if value is None:
            return None
        if isinstance(value, list):
            value = "|".join(f"(?:{v})" for v in value)
        return re.compile(value)

    name_pattern = compile_filter(args.name_pattern)
    memo_filter = compile_filter(args.memo_filter)

    # 各チャンネルからファイルを集める
    files: list = []
    for ch_spec in channels:
        try:
            ch_id = resolve_channel_id(client, ch_spec)
        except SystemExit as e:
            print(f"[警告] チャンネル '{ch_spec}' をスキップ: {e}", file=sys.stderr)
            continue

        print(f"\n=== チャンネル {ch_spec} ({ch_id}) を処理 ===", file=sys.stderr)
        try:
            messages = fetch_messages(
                client, ch_id, args.limit, since_ts, include_threads=args.threads
            )
        except SlackApiError as e:
            print(f"[警告] {ch_spec}: Slack API error: {e.response['error']}",
                  file=sys.stderr)
            continue

        if args.debug:
            print(f"=== DEBUG: {ch_spec} 取得メッセージ {len(messages)} 件 ===",
                  file=sys.stderr)

        # tracking+teamsモード: Teams URLを含む親メッセージを収集 (Excel不要)
        # それ以外: ファイルを収集 (従来挙動)
        if args.extract == "tracking" and args.tracking_source == "teams":
            ch_files = collect_teams_posts(messages, target_date=target_date)
            kind_label = "Teams投稿"
        else:
            ch_files = collect_files(messages, client, target_date=target_date,
                                     name_pattern=name_pattern)
            kind_label = "ファイル"
        # 各アイテムにチャンネルIDを紐付け
        for f in ch_files:
            f["_channel_id"] = ch_id
            f["_channel_spec"] = ch_spec
        print(f"  → {len(ch_files)} 件の{kind_label}候補", file=sys.stderr)
        files.extend(ch_files)

    if not files:
        if args.extract == "tracking" and args.tracking_source == "teams":
            print("Teams URL を含む投稿が全チャンネルで見つかりませんでした。",
                  file=sys.stderr)
        else:
            print("対象ファイル (.xlsx/.xls/.docx/.csv) は全チャンネルで見つかりませんでした。",
                  file=sys.stderr)
        return

    if args.extract == "tracking" and args.tracking_source == "teams":
        print(f"\n全チャンネル合計 {len(files)} 件のTeams投稿を処理します...",
              file=sys.stderr)
    else:
        print(f"\n全チャンネル合計 {len(files)} 件のファイルを処理します...",
              file=sys.stderr)

    # SharePoint URLが含まれているなら Graph API トークンを取得
    needs_graph = any(f.get("_sharepoint_url") for f in files) \
        or (args.extract == "tracking" and args.tracking_source == "teams")
    extra_scopes = []
    if args.extract == "tracking" and args.tracking_source == "teams":
        extra_scopes.append("OnlineMeetings.Read")
        extra_scopes.append("Calendars.Read")
        # ドライバー予定表を指定している場合は共有読取スコープも追加
        if getattr(args, "track_calendars", None):
            extra_scopes.append("Calendars.Read.Shared")
    graph_token = get_graph_token(extra_scopes=extra_scopes) if needs_graph else None

    # ドライバー予定表のプリフェッチ (Teamsモードでtrack_calendarsが指定されている場合)
    track_event_map: dict = {}
    if (args.extract == "tracking" and args.tracking_source == "teams"
            and graph_token and getattr(args, "track_calendars", None)):
        drivers = args.track_calendars
        if isinstance(drivers, str):
            drivers = [drivers]
        print(f"\n=== ドライバー予定表をプリフェッチ ({len(drivers)}人) ===",
              file=sys.stderr)
        track_event_map = fetch_track_calendars(
            drivers, graph_token, target_date, days_window=14
        )
        print(f"[track-cal] 合計 {len(track_event_map)} 件のオンライン会議を "
              f"JoinUrl で索引化", file=sys.stderr)

    def fetch_file_bytes(f: dict) -> tuple[str, bytes]:
        """ファイルエントリを判別してダウンロード。(filename, bytes) を返す。"""
        if f.get("_sharepoint_url"):
            real_name, data = download_sharepoint_file(f["_sharepoint_url"], graph_token)
            if real_name:
                f["name"] = real_name
            return f.get("name", "sharepoint_file"), data
        return f.get("name", "unknown"), download_file(f, token)

    # SharePointファイルは事前にメタデータだけ取得して、対象外なら除外する
    if needs_graph:
        import base64
        headers = {"Authorization": f"Bearer {graph_token}"}
        validated = []
        default_year = target_date.year if target_date else datetime.now().year
        for f in files:
            sp_url = f.get("_sharepoint_url")
            if not sp_url:
                validated.append(f)
                continue
            try:
                encoded = base64.urlsafe_b64encode(sp_url.encode()).decode().rstrip("=")
                meta_resp = requests.get(
                    f"https://graph.microsoft.com/v1.0/shares/u!{encoded}/driveItem",
                    headers=headers, timeout=30,
                )
                if not meta_resp.ok:
                    if args.debug:
                        print(f"[debug] SP メタ取得失敗 ({meta_resp.status_code}): {sp_url[:80]}",
                              file=sys.stderr)
                    continue
                real_name = meta_resp.json().get("name", "")
                real_ext = Path(real_name).suffix.lower()
                if real_ext not in TARGET_EXTS:
                    if args.debug:
                        print(f"[debug] SP 拡張子対象外 ({real_name}): {sp_url[:80]}",
                              file=sys.stderr)
                    continue
                # ファイル名パターンチェック (SharePointは実名でここで判定)
                if name_pattern and not name_pattern.search(real_name):
                    if args.debug:
                        print(f"[debug] SP 名前パターン不一致 ({real_name})",
                              file=sys.stderr)
                    continue
                # 日付フィルタ: 実ファイル名でも確認
                if target_date is not None:
                    file_dates = extract_dates(real_name, default_year)
                    msg_dates = f.get("_msg_dates", set())
                    if target_date not in (msg_dates | file_dates):
                        if args.debug:
                            print(f"[debug] SP 日付不一致 ({real_name}): {sp_url[:80]}",
                                  file=sys.stderr)
                        continue
                f["name"] = real_name
                validated.append(f)
            except Exception as e:
                if args.debug:
                    print(f"[debug] SP 検証エラー: {type(e).__name__}: {e}", file=sys.stderr)
                continue
        files = validated

    if not files:
        print("対象ファイルが0件になりました(SharePointメタ検証後)。", file=sys.stderr)
        return

    print(f"検証後 {len(files)} 件を処理します...", file=sys.stderr)

    if args.extract == "driving":
        import json
        parser_fn, allowed_exts = STRUCTURED_EXTRACTORS["driving"]
        records = []
        total = len(files)
        for idx, f in enumerate(files, 1):
            posted = (datetime.fromtimestamp(float(f["_posted_ts"])).isoformat(timespec="seconds")
                      if f.get("_posted_ts") else "")
            print(f"  [{idx}/{total}] {f.get('name', '?')[:60]}", file=sys.stderr)
            try:
                name, data = fetch_file_bytes(f)
                ext = Path(name).suffix.lower()
                if ext not in allowed_exts:
                    continue
                parsed = parser_fn(data)

                # メモフィルタ適用: 各日のentriesから、memoがマッチする行のみ残す
                if memo_filter:
                    filtered_days = []
                    for day in parsed.get("days", []):
                        kept = [e for e in day.get("entries", [])
                                if memo_filter.search(e.get("memo", ""))]
                        if kept:
                            new_day = {k: v for k, v in day.items() if k != "entries"}
                            new_day["entries"] = kept
                            filtered_days.append(new_day)
                    parsed = {"days": filtered_days}
                    if not filtered_days:
                        continue  # マッチ無しの場合はファイルごと省略

                records.append({"file": name, "posted": posted, "user": f.get("user", ""),
                                "data": parsed})
            except Exception as e:
                records.append({"file": f.get("name", "?"), "posted": posted,
                                "user": f.get("user", ""),
                                "error": f"{type(e).__name__}: {e}"})
        result = json.dumps(records, ensure_ascii=False, indent=2)
    elif args.extract == "tracking":
        import json
        parser_fn, allowed_exts = STRUCTURED_EXTRACTORS["driving"]
        user_cache: dict = {}

        # appendモード: 既存ファイルを読み込んでURLセットを作る
        existing_records: list = []
        existing_labeled: list = []
        existing_urls: set = set()
        if args.append and args.out and os.path.exists(args.out):
            try:
                with open(args.out, encoding="utf-8") as f:
                    existing = json.load(f)
                for item in existing:
                    if isinstance(item, dict) and "__labeled_view__" in item:
                        existing_labeled = item["__labeled_view__"]
                    elif isinstance(item, list) and len(item) >= 4 \
                            and isinstance(item[3], dict):
                        existing_records.append(item)
                        url = item[3].get("url", "")
                        if url:
                            existing_urls.add(url)
                print(f"[append] 既存 {len(existing_records)} 件を読み込み "
                      f"(うちURL付き {len(existing_urls)} 件)", file=sys.stderr)
            except Exception as e:
                print(f"[警告] 既存result読み込み失敗、新規作成として進行: {e}",
                      file=sys.stderr)

        records: list = []
        labeled_records: list = []
        legs_records: list = []
        # URL→投稿日ISO のマップ (logs.json 構築用)
        url_to_posted_iso: dict = {}
        skipped_count = 0
        total = len(files)
        default_year = target_date.year if target_date else datetime.now().year
        for idx, f in enumerate(files, 1):
            print(f"  [{idx}/{total}] {f.get('name', '?')[:60]}", file=sys.stderr)
            msg = f.get("_msg") or {}
            posted_iso = (
                datetime.fromtimestamp(float(f["_posted_ts"])).isoformat(timespec="seconds")
                if f.get("_posted_ts") else ""
            )
            try:
                meta = extract_tracking_metadata(msg, client, user_cache, default_year)
                if args.debug:
                    src = msg.get("_parent_msg") or msg
                    snippet = (src.get("text") or "").replace("\n", " | ")[:200]
                    print(f"    [tracking] 親本文先頭: {snippet}", file=sys.stderr)
                    meta_clean = {k: v for k, v in meta.items() if not k.startswith("_")}
                    print(f"    [tracking] 抽出メタ: {meta_clean}", file=sys.stderr)

                # tracking_source=teams: Teams会議から件名・時刻を取得して上書き
                teams_time_range = None
                if args.tracking_source == "teams":
                    if not graph_token:
                        print(f"    [teams診断] graph_token が None。"
                              f"Teams API呼び出し不可", file=sys.stderr)
                    else:
                        src = msg.get("_parent_msg") or msg
                        teams_urls = extract_teams_urls(src)
                        if not teams_urls:
                            print(f"    [teams診断] このメッセージに Teams URL なし "
                                  f"(text長={len(src.get('text') or '')})",
                                  file=sys.stderr)
                        else:
                            meeting = get_teams_meeting(teams_urls[0], graph_token)
                            # track_event_map をフォールバック (Teams API取れなかった場合)
                            if not meeting and track_event_map:
                                url_in = teams_urls[0]
                                if url_in in track_event_map:
                                    ev = track_event_map[url_in]
                                    s_str = ev["start"]["dateTime"]
                                    e_str = ev["end"]["dateTime"]
                                    s_str = (s_str.split(".")[0] + "Z"
                                             if "." in s_str
                                             else (s_str if s_str.endswith("Z")
                                                   else s_str + "Z"))
                                    e_str = (e_str.split(".")[0] + "Z"
                                             if "." in e_str
                                             else (e_str if e_str.endswith("Z")
                                                   else e_str + "Z"))
                                    meeting = {
                                        "subject": ev.get("subject", ""),
                                        "startDateTime": s_str,
                                        "endDateTime": e_str,
                                        "_source": "driverCalendar",
                                    }
                                    print(f"    [track-cal] マッチ: "
                                          f"{ev.get('subject', '')[:60]}",
                                          file=sys.stderr)
                            if meeting:
                                subject = meeting.get("subject", "") or ""
                                tmeta = parse_teams_subject(subject)
                                # Trackname / Track-num / Customer のみTeams由来で上書き
                                # Route は Slack本文の `自動運転区間：` から取得する従来挙動を維持
                                for k in ("Trackname", "Track-num", "Customer"):
                                    v = tmeta.get(k)
                                    if v:
                                        meta[k] = v
                                # Time-range = Teams開始/終了時刻 (JST)
                                from datetime import timezone, timedelta
                                JST = timezone(timedelta(hours=9))
                                try:
                                    slack_ts = float(f.get("_posted_ts") or 0)
                                    if slack_ts > 0:
                                        posted_jst = datetime.fromtimestamp(slack_ts, tz=JST)
                                        start_str = meeting.get("startDateTime", "")
                                        end_str = meeting.get("endDateTime", "")
                                        if start_str and end_str:
                                            s_utc = datetime.fromisoformat(
                                                start_str.replace("Z", "+00:00")
                                            )
                                            e_utc = datetime.fromisoformat(
                                                end_str.replace("Z", "+00:00")
                                            )
                                            s_dt = s_utc.astimezone(JST)
                                            e_dt = e_utc.astimezone(JST)
                                            meta["_teams_start_dt"] = s_dt
                                            meta["_teams_end_dt"] = e_dt
                                            # 日付はTeams会議の開始日を採用
                                            meta["_title_date"] = s_dt.date()
                                            date_str = s_dt.strftime("%Y/%m/%d")
                                            teams_time_range = (
                                                f"{date_str}T{s_dt.strftime('%H:%M')}+9:00/"
                                                f"{date_str}T{e_dt.strftime('%H:%M')}+9:00"
                                            )
                                            meta["_teams_time_range"] = teams_time_range
                                            track_disp = tmeta.get("Trackname") or "?"
                                            print(
                                                f"    [teams時刻] {track_disp}: "
                                                f"Teams開始(UTC)={start_str} → "
                                                f"JST {s_dt.strftime('%Y-%m-%d %H:%M')}",
                                                file=sys.stderr,
                                            )
                                            print(
                                                f"    [teams時刻] {track_disp}: "
                                                f"Teams終了(UTC)={end_str} → "
                                                f"JST {e_dt.strftime('%Y-%m-%d %H:%M')}",
                                                file=sys.stderr,
                                            )
                                            print(
                                                f"    [teams時刻] {track_disp}: "
                                                f"Slack投稿日={posted_jst.strftime('%Y-%m-%d')} / "
                                                f"Teams会議日={s_dt.date()} "
                                                f"→ 採用={date_str}",
                                                file=sys.stderr,
                                            )
                                            print(
                                                f"    [teams時刻] {track_disp}: "
                                                f"最終 Time-range = {teams_time_range}",
                                                file=sys.stderr,
                                            )
                                        else:
                                            print(f"    [teams診断] meeting応答に "
                                                  f"startDateTime/endDateTime が無い: "
                                                  f"{list(meeting.keys())}",
                                                  file=sys.stderr)
                                except (ValueError, KeyError) as e:
                                    print(f"    [teams] 時刻パース失敗: {e}",
                                          file=sys.stderr)

                # 厳密な日付フィルタ: target_date と _title_date が一致しないものは除外
                # (collect_files の日付フィルタは "本文に出現すれば通過" と緩いため、
                #  最終的にここで厳しくチェック)
                if target_date is not None:
                    rec_date = meta.get("_title_date")
                    if rec_date != target_date:
                        if args.debug:
                            print(f"    [filter] 日付不一致でスキップ: "
                                  f"レコード日付={rec_date}, target={target_date}",
                                  file=sys.stderr)
                        continue

                url = get_message_permalink(client, f.get("_channel_id", ""),
                                            f["_posted_ts"]) \
                    if f.get("_posted_ts") and f.get("_channel_id") else ""

                # 既存にあればここでスキップ
                if args.append and url and url in existing_urls:
                    skipped_count += 1
                    if args.debug:
                        print(f"    [append] スキップ(既存): {url}", file=sys.stderr)
                    continue

                # Excelの読み込みは tracking_source=excel のときだけ
                day_match = None
                if args.tracking_source == "excel":
                    name, data = fetch_file_bytes(f)
                    ext = Path(name).suffix.lower()
                    if ext not in allowed_exts:
                        if args.debug:
                            print(f"[debug] tracking: {ext} は未対応", file=sys.stderr)
                        continue
                    parsed = parser_fn(data)

                    title_date = meta.get("_title_date")
                    if title_date:
                        iso = title_date.strftime("%Y-%m-%d")
                        for d in parsed.get("days", []):
                            if d.get("date") == iso and (d.get("ad_start") or d.get("ad_end")):
                                day_match = d
                                break
                        if day_match is None:
                            for d in parsed.get("days", []):
                                if d.get("date") == iso:
                                    day_match = d
                                    break

                records.append(build_tracking_record(meta, day_match, posted_iso, url))
                labeled_records.append(build_labeled_record(meta, day_match, url))
                legs_records.append(build_legs_record(meta, day_match, url))
                if url:
                    existing_urls.add(url)
                    if posted_iso:
                        url_to_posted_iso[url] = posted_iso
            except Exception as e:
                if args.debug:
                    print(f"[debug] tracking エラー ({f.get('name')}): "
                          f"{type(e).__name__}: {e}", file=sys.stderr)
                continue

        # ── レコード日付フィルタ: target_date と一致するレコードだけ残す ──
        if target_date:
            filtered_records, filtered_labeled, filtered_legs = [], [], []
            for rec, lrec, lgrec in zip(records, labeled_records, legs_records):
                rec_date = None
                date_part = (rec[1] or "").split("|", 1)
                if len(date_part) == 2:
                    try:
                        rec_date = datetime.strptime(date_part[1], "%Y/%m/%d").date()
                    except ValueError:
                        pass
                if rec_date is None or rec_date == target_date:
                    filtered_records.append(rec)
                    filtered_labeled.append(lrec)
                    filtered_legs.append(lgrec)
                elif args.debug:
                    print(f"    [date-filter] スキップ {rec[0]}: {rec_date} != "
                          f"{target_date}", file=sys.stderr)
            dropped = len(records) - len(filtered_records)
            if dropped > 0:
                print(f"[date-filter] target {target_date} 以外を "
                      f"{dropped} 件除外", file=sys.stderr)
            records = filtered_records
            labeled_records = filtered_labeled
            legs_records = filtered_legs

        # 統計表示
        if args.append:
            print(f"[append] 新規 {len(records)} 件、重複スキップ {skipped_count} 件",
                  file=sys.stderr)

        # 既存 + 新規 をマージ
        final_records = existing_records + records
        final_labeled = existing_labeled + labeled_records
        combined: list = list(final_records)
        combined.append({"__labeled_view__": final_labeled})
        result = json.dumps(combined, ensure_ascii=False, indent=2)
    else:
        chunks = []
        total = len(files)
        for idx, f in enumerate(files, 1):
            posted = (datetime.fromtimestamp(float(f["_posted_ts"])).isoformat(timespec="seconds")
                      if f.get("_posted_ts") else "")
            user = f.get("user", "")
            print(f"  [{idx}/{total}] {f.get('name', '?')[:60]}", file=sys.stderr)
            try:
                name, data = fetch_file_bytes(f)
                ext = Path(name).suffix.lower()
                text = EXTRACTORS[ext](data) if ext in EXTRACTORS else \
                    f"[拡張子 {ext} は未対応]"
            except Exception as e:
                name = f.get("name", "?")
                text = f"[読み込みエラー: {type(e).__name__}: {e}]"
            header = (
                f"\n{'=' * 70}\n"
                f"FILE   : {name}\n"
                f"POSTED : {posted}   BY: {user}\n"
                f"{'=' * 70}"
            )
            chunks.append(header + "\n" + text)
        result = "\n".join(chunks)
    if args.out:
        Path(args.out).write_text(result, encoding="utf-8")
        print(f"結果を {args.out} に保存しました。", file=sys.stderr)
    else:
        print(result)

    # legs.json への保存 (trackingモードのみ)
    legs_new_count = legs_skipped_count = legs_total_count = 0
    if args.legs_out and args.extract == "tracking":
        try:
            existing_legs: list = []
            existing_keys: set = set()
            if os.path.exists(args.legs_out):
                # BOMを見て適切なエンコーディングで読む
                with open(args.legs_out, "rb") as f:
                    raw = f.read()
                if raw.startswith(b'\xff\xfe'):
                    content = raw.decode('utf-16-le', errors='replace').strip()
                    if content.startswith('\ufeff'):
                        content = content[1:]
                    print(f"[legs] {args.legs_out} はUTF-16 LE形式。"
                          f"UTF-8で再保存します", file=sys.stderr)
                elif raw.startswith(b'\xfe\xff'):
                    content = raw.decode('utf-16-be', errors='replace').strip()
                    if content.startswith('\ufeff'):
                        content = content[1:]
                    print(f"[legs] {args.legs_out} はUTF-16 BE形式。"
                          f"UTF-8で再保存します", file=sys.stderr)
                elif raw.startswith(b'\xef\xbb\xbf'):
                    content = raw[3:].decode('utf-8', errors='replace').strip()
                else:
                    content = raw.decode('utf-8', errors='replace').strip()

                if not content:
                    print(f"[legs] {args.legs_out} は空。新規作成扱いで進めます",
                          file=sys.stderr)
                else:
                    try:
                        existing_legs = json.loads(content)
                        if not isinstance(existing_legs, list):
                            print(f"[警告] {args.legs_out} の中身が配列でない。"
                                  f"新規作成扱いに切替", file=sys.stderr)
                            existing_legs = []
                    except json.JSONDecodeError as e:
                        print(f"[警告] {args.legs_out} のJSONが壊れている "
                              f"({e})。新規作成扱いに切替", file=sys.stderr)
                        existing_legs = []
                for rec in existing_legs:
                    if isinstance(rec, dict):
                        k = legs_dedup_key(rec)
                        if any(k):
                            existing_keys.add(k)
                if existing_legs:
                    print(f"[legs] 既存 {len(existing_legs)} 件を読み込み",
                          file=sys.stderr)

            new_legs: list = []
            for rec in legs_records:
                k = legs_dedup_key(rec)
                if any(k) and k in existing_keys:
                    legs_skipped_count += 1
                    if args.debug:
                        print(f"    [legs] スキップ(既存): {k}", file=sys.stderr)
                    continue
                new_legs.append(rec)
                if any(k):
                    existing_keys.add(k)

            legs_new_count = len(new_legs)
            final_legs = existing_legs + new_legs
            legs_total_count = len(final_legs)
            Path(args.legs_out).write_text(
                json.dumps(final_legs, ensure_ascii=False, indent=2),
                encoding="utf-8"
            )
            print(f"[legs] {args.legs_out} に保存 "
                  f"(新規 {legs_new_count} 件 / 重複スキップ {legs_skipped_count} 件 / "
                  f"累計 {legs_total_count} 件)", file=sys.stderr)
        except Exception as e:
            print(f"[警告] legs.json 出力エラー: {e}", file=sys.stderr)

    # logs.json 出力 (legs.jsonに書き込んだ投稿の投稿日履歴)
    if args.logs_out and args.extract == "tracking":
        try:
            # 既存ログのURL重複チェック用集合
            existing_log_urls = {
                e.get("url") for e in existing_logs
                if isinstance(e, dict) and e.get("url")
            }
            # 今回処理した legs_records を対象に、URL未登録のものを追記
            new_log_entries: list = []
            for rec in legs_records:
                if not (isinstance(rec, list) and len(rec) >= 4):
                    continue
                meta_dict = rec[3] if isinstance(rec[3], dict) else {}
                url = meta_dict.get("url", "")
                trackname = rec[0] if rec else ""
                posted_at = url_to_posted_iso.get(url, "")
                if not url or url in existing_log_urls:
                    continue
                if not posted_at:
                    continue
                new_log_entries.append({
                    "trackname": trackname,
                    "url": url,
                    "posted_at": posted_at,
                })
                existing_log_urls.add(url)

            final_logs = existing_logs + new_log_entries
            Path(args.logs_out).write_text(
                json.dumps(final_logs, ensure_ascii=False, indent=2),
                encoding="utf-8"
            )
            print(f"[logs] {args.logs_out} に保存 "
                  f"(新規 {len(new_log_entries)} 件 / 累計 {len(final_logs)} 件)",
                  file=sys.stderr)
        except Exception as e:
            print(f"[警告] logs.json 出力エラー: {e}", file=sys.stderr)

    # Slack通知 (Incoming Webhook)
    if args.notify_webhook_url:
        # legs_records の完全性チェック (trackingモードのみ)
        status_line = "✅ success: 運行記録の読み込みが完了しました"
        incomplete_rec_indices: set = set()
        if args.extract == "tracking":
            try:
                if legs_records:
                    for i, lrec in enumerate(legs_records):
                        ok, _ = is_legs_record_complete(lrec)
                        if not ok:
                            incomplete_rec_indices.add(i)
                    if incomplete_rec_indices:
                        status_line = (
                            f"❌ failed: 運行記録に不完全なレコードがあります "
                            f"({len(incomplete_rec_indices)}/{len(legs_records)} 件)"
                        )
                else:
                    # 0件の場合は success (失敗ではなく対象なし扱い)
                    status_line = "✅ success: (今回追加されたレコードはありません)"
            except NameError:
                pass

        lines = [status_line]
        if target_date:
            lines.append(f"対象日: {target_date.isoformat()}")

        if args.extract == "tracking":
            try:
                shown = records if records else []
                MAX_IN_NOTIF = 30
                for i, rec in enumerate(shown[:MAX_IN_NOTIF]):
                    track = rec[0] or "(Trackname不明)"
                    meta = rec[3] if isinstance(rec[3], dict) else {}
                    customer = meta.get("loaded liggage", "")
                    section = meta.get("selfdrive section", "")
                    url = meta.get("url", "")

                    # 対応する legs_record の完全性を見て印を付ける
                    marker = ""
                    if i in incomplete_rec_indices:
                        _, missing = is_legs_record_complete(legs_records[i])
                        marker = f" ⚠️ 未取得: {', '.join(missing)}"

                    lines.append("")
                    lines.append(f"{track}{marker}")
                    lines.append(f'運行名: "{customer}"')
                    lines.append(f'自動運転区間: "{section}"')
                    if url:
                        lines.append(f'URL: <{url}|スレッドを開く>')
                if len(shown) > MAX_IN_NOTIF:
                    lines.append("")
                    lines.append(f"…他 {len(shown) - MAX_IN_NOTIF} 件")
            except NameError:
                pass

        send_slack_notification(args.notify_webhook_url, "\n".join(lines))


if __name__ == "__main__":
    main()